import os
import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

# Load questions from JSON
json_path = os.path.join(os.path.dirname(__file__), 'questions.json')
with open(json_path, 'r', encoding='utf-8') as f:
    questions = json.load(f)

# Sort questions by ID descending
questions.sort(key=lambda x: x.get('id', 0), reverse=True)

# Helper to clean HTML tags
def clean_html(text):
    if not text:
        return ""
    # Replace some basic paragraph elements with newlines if needed, otherwise strip all
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'<p[^>]*>', '', text)
    text = re.sub(r'</p>', '\n', text)
    text = re.sub(r'<[^>]+>', '', text)
    # Decode HTML entities
    text = text.replace('&nbsp;', ' ').replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&').replace('&quot;', '"').replace('&apos;', "'")
    # Clean multiple newlines
    text = re.sub(r'\n+', '\n', text)
    return text.strip()

# Helper for difficulty stars
def get_diff_stars(difficult):
    if not difficult:
        return "⭐"
    return "⭐" * min(max(difficult, 1), 5)

# Helper for question type names
TYPE_NAMES = {
    1: "单选题",
    2: "多选题",
    3: "判断题",
    4: "填空题",
    5: "简答题"
}

# Group questions by type
grouped = {}
for q in questions:
    q_type = q.get('questionType', 1)
    grouped.setdefault(q_type, []).append(q)

# ----------------- Excel Generation -----------------
def generate_excel():
    wb = Workbook()
    ws = wb.active
    ws.title = "题库汇总"
    
    # Enable grid lines
    ws.views.sheetView[0].showGridLines = True
    
    # Headers
    headers = ["ID", "题型", "题干", "选项 (A/B/C/D...)", "正确答案", "难度", "分值", "制度来源", "题库标签", "解析"]
    ws.append(headers)
    
    # Style definitions
    header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid") # Classic Navy Blue
    alignment_center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    alignment_left = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin_border = Border(
        left=Side(style='thin', color='D9D9D9'),
        right=Side(style='thin', color='D9D9D9'),
        top=Side(style='thin', color='D9D9D9'),
        bottom=Side(style='thin', color='D9D9D9')
    )
    
    # Style header row
    for col_idx in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = alignment_center
        cell.border = thin_border
    
    row_num = 2
    # Add data grouped by type for readability
    for q_type in sorted(grouped.keys()):
        for q in grouped[q_type]:
            # Format choices
            choices = []
            for item in q.get('items', []):
                prefix = item.get('prefix', '')
                content = clean_html(item.get('content', ''))
                choices.append(f"{prefix}. {content}")
            choices_str = "\n".join(choices)
            
            # Format correct answers
            correct = q.get('correct', '')
            if not correct and q.get('correctArray'):
                correct = ", ".join(q.get('correctArray', []))
            
            ws.append([
                q.get('id'),
                TYPE_NAMES.get(q_type, "其他"),
                clean_html(q.get('title', '')),
                choices_str,
                correct,
                get_diff_stars(q.get('difficult', 1)),
                f"{q.get('score', '2')}分",
                q.get('regulationSource', ''),
                q.get('tag', ''),
                clean_html(q.get('analyze', ''))
            ])
            
            # Style data row
            for col_idx in range(1, len(headers) + 1):
                cell = ws.cell(row=row_num, column=col_idx)
                cell.font = Font(name="微软雅黑", size=10)
                cell.border = thin_border
                
                # Alignments
                if col_idx in [1, 2, 5, 6, 7, 8, 9]:
                    cell.alignment = alignment_center
                else:
                    cell.alignment = alignment_left
            
            row_num += 1
            
    # Set column widths
    ws.column_dimensions['A'].width = 8
    ws.column_dimensions['B'].width = 12
    ws.column_dimensions['C'].width = 40
    ws.column_dimensions['D'].width = 35
    ws.column_dimensions['E'].width = 12
    ws.column_dimensions['F'].width = 12
    ws.column_dimensions['G'].width = 8
    ws.column_dimensions['H'].width = 25
    ws.column_dimensions['I'].width = 15
    ws.column_dimensions['J'].width = 40

    # Auto height adjust for rows
    for r in range(1, row_num):
        ws.row_dimensions[r].height = None # Excel will auto height based on text wrapping
        
    excel_file = os.path.join(os.path.dirname(__file__), 'question_bank.xlsx')
    wb.save(excel_file)
    print(f"Excel generated at: {excel_file}")

# ----------------- Word Generation -----------------
def generate_word():
    doc = Document()
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("XZS 考试系统 - 完整导出题库")
    title_run.font.name = "微软雅黑"
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(31, 78, 120)
    
    # Subtitle
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run(f"总题目数: {len(questions)} 道 | 自动导出日期: 2026年5月")
    subtitle_run.font.name = "微软雅黑"
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph() # Spacer
    
    # Format each group
    type_index_map = {1: "一、单选题", 2: "二、多选题", 3: "三、判断题", 4: "四、填空题", 5: "五、简答题"}
    
    for q_type in sorted(type_index_map.keys()):
        if q_type not in grouped or not grouped[q_type]:
            continue
            
        # Section Header
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(type_index_map[q_type])
        h_run.font.name = "微软雅黑"
        h_run.font.size = Pt(14)
        h_run.font.bold = True
        h_run.font.color.rgb = RGBColor(31, 78, 120)
        
        # Add questions in this group
        for idx, q in enumerate(grouped[q_type], 1):
            qp = doc.add_paragraph()
            # Question text
            title_text = clean_html(q.get('title', ''))
            score = q.get('score', '2')
            q_run = qp.add_run(f"{idx}. [ID:{q.get('id')}] {title_text} （{score}分）")
            q_run.font.name = "微软雅黑"
            q_run.font.size = Pt(10.5)
            q_run.font.bold = True
            
            # Format choices (single/multi-choice)
            if q.get('items'):
                # Check if we should display options inline or vertically based on length
                choices = []
                for item in q.get('items', []):
                    prefix = item.get('prefix', '')
                    content = clean_html(item.get('content', ''))
                    choices.append(f"{prefix}. {content}")
                
                # Check total options character length
                total_len = sum(len(c) for c in choices)
                op_p = doc.add_paragraph()
                op_p.paragraph_format.left_indent = Pt(18)
                op_p.paragraph_format.line_spacing = 1.15
                
                if total_len < 50:
                    # Render in single line
                    op_run = op_p.add_run("    ".join(choices))
                else:
                    # Render one choice per line
                    op_run = op_p.add_run("\n".join(choices))
                    
                op_run.font.name = "微软雅黑"
                op_run.font.size = Pt(10)
                op_run.font.color.rgb = RGBColor(64, 64, 64)
            
            # Answer & Meta paragraph (Tightly packed)
            correct = q.get('correct', '')
            if not correct and q.get('correctArray'):
                correct = ", ".join(q.get('correctArray', []))
            
            meta_p = doc.add_paragraph()
            meta_p.paragraph_format.left_indent = Pt(18)
            meta_p.paragraph_format.space_after = Pt(12) # Gap before next question
            
            # Correct Answer
            ans_run = meta_p.add_run("【正确答案】")
            ans_run.font.name = "微软雅黑"
            ans_run.font.size = Pt(9.5)
            ans_run.font.bold = True
            ans_run.font.color.rgb = RGBColor(46, 117, 89) # Deep green
            
            ans_val = meta_p.add_run(f" {correct}    ")
            ans_val.font.name = "微软雅黑"
            ans_val.font.size = Pt(9.5)
            ans_val.font.bold = True
            ans_val.font.color.rgb = RGBColor(0, 0, 0)
            
            # Difficulty
            diff_run = meta_p.add_run("【难度】")
            diff_run.font.name = "微软雅黑"
            diff_run.font.size = Pt(9.5)
            diff_run.font.bold = True
            
            diff_val = meta_p.add_run(f" {get_diff_stars(q.get('difficult', 1))}    ")
            diff_val.font.name = "微软雅黑"
            diff_val.font.size = Pt(9.5)
            
            # Source
            if q.get('regulationSource'):
                src_run = meta_p.add_run("【制度来源】")
                src_run.font.name = "微软雅黑"
                src_run.font.size = Pt(9.5)
                src_run.font.bold = True
                
                src_val = meta_p.add_run(f" {q.get('regulationSource')}    ")
                src_val.font.name = "微软雅黑"
                src_val.font.size = Pt(9.5)
                src_val.font.color.rgb = RGBColor(128, 0, 0)
                
            # Tag
            if q.get('tag'):
                tag_run = meta_p.add_run("【标签】")
                tag_run.font.name = "微软雅黑"
                tag_run.font.size = Pt(9.5)
                tag_run.font.bold = True
                
                tag_val = meta_p.add_run(f" {q.get('tag')}    ")
                tag_val.font.name = "微软雅黑"
                tag_val.font.size = Pt(9.5)
            
            # Analysis
            if q.get('analyze'):
                meta_p.add_run("\n") # Line break for analysis to keep things clean
                ana_run = meta_p.add_run("【解析】")
                ana_run.font.name = "微软雅黑"
                ana_run.font.size = Pt(9.5)
                ana_run.font.bold = True
                ana_run.font.color.rgb = RGBColor(128, 128, 128)
                
                ana_val = meta_p.add_run(f" {clean_html(q.get('analyze', ''))}")
                ana_val.font.name = "微软雅黑"
                ana_val.font.size = Pt(9.5)
                ana_val.font.color.rgb = RGBColor(100, 100, 100)
                
    word_file = os.path.join(os.path.dirname(__file__), 'question_bank.docx')
    doc.save(word_file)
    print(f"Word document generated at: {word_file}")

if __name__ == '__main__':
    generate_excel()
    generate_word()
